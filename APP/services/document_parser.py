"""
Document Parser Service
PDF, HWP, DOCX, 이미지 파일 파싱을 처리하는 서비스
EasyOCR 기반으로 최적화됨
"""

import os
import logging
from typing import Dict, Any, Optional, List
from abc import ABC, abstractmethod
from pathlib import Path
import mimetypes
import re

# PDF 파싱
import fitz  # PyMuPDF
import pdfplumber
from PyPDF2 import PdfReader

# DOCX 파싱
from docx import Document as DocxDocument

# HWP 파싱
import olefile
import zipfile
import xml.etree.ElementTree as ET

# 이미지/OCR (EasyOCR)
import numpy as np
from PIL import Image
import easyocr

logger = logging.getLogger(__name__)

# 🔹 전역 OCR Reader (한 번만 로드)
_OCR_READER = None

def get_ocr_reader():
    """EasyOCR Reader를 한 번만 로드하여 재사용"""
    global _OCR_READER
    if _OCR_READER is None:
        logger.info("Initializing EasyOCR Reader...")
        _OCR_READER = easyocr.Reader(['ko', 'en'], gpu=False)
    return _OCR_READER


def clean_pdf_text(text: str) -> str:
    """
    PDF에서 추출한 텍스트에서 불필요한 메타데이터 제거
    - 폰트 인코딩 데이터 (Base64)
    - 각주 번호만 있는 줄
    - PDF 메타데이터
    """
    if not text:
        return ""
    
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        # 빈 줄은 유지
        if not stripped:
            cleaned_lines.append('')
            continue
        
        # Base64로 보이는 긴 문자열이 80% 이상인 줄 제거
        if len(stripped) > 20:
            base64_chars = len(re.findall(r'[A-Za-z0-9+/=]', stripped))
            if base64_chars / len(stripped) > 0.8:
                continue
        
        # 각주 번호만 있는 줄 (^1, ^2, (^3) 등)
        if re.match(r'^\(?\^?\d+[\.\)]\)?\s*$', stripped):
            continue
        
        # PDF 메타데이터 키워드만 있는 줄
        if stripped in ['SHA1', 'MD5', '{}', '[]', '()', 'IAA=']:
            continue
        
        # "--- Page N ---" 구분선은 유지
        if stripped.startswith('--- Page'):
            cleaned_lines.append(line)
            continue
        
        # 나머지는 유지
        cleaned_lines.append(line)
    
    text = '\n'.join(cleaned_lines)
    
    # 연속된 빈 줄 정리 (3줄 이상 → 2줄)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


class DocumentParser(ABC):
    """문서 파서 추상 클래스"""
    
    @abstractmethod
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """
        파일을 파싱하여 텍스트와 메타데이터를 추출
        
        Args:
            file_path: 파싱할 파일 경로
            
        Returns:
            {
                "text": str,           # 추출된 텍스트
                "page_count": int,     # 페이지 수
                "has_tables": bool,    # 테이블 포함 여부
                "confidence": float,   # 신뢰도 (OCR의 경우)
                "metadata": dict       # 추가 메타데이터
            }
        """
        pass
    
    def clean_text(self, text: str) -> str:
        """텍스트 정리"""
        # 연속된 공백을 하나로
        text = re.sub(r'\s+', ' ', text)
        
        # 연속된 줄바꿈 정리 (3개 이상 → 2개)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # 양쪽 공백 제거
        text = text.strip()
        
        return text


class PDFParser(DocumentParser):
    """PDF 파일 파서 - PyMuPDF + pdfplumber 조합"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """PDF 파싱 - 벡터 텍스트 우선 → OCR 폴백"""
        try:
            logger.info(f"Parsing PDF file: {file_path}")
            
            # ---------------- 1. 벡터 텍스트 추출 (PyMuPDF) ----------------
            doc = fitz.open(file_path)
            text_parts = []
            has_tables = False
            page_count = len(doc)
            
            for page_num, page in enumerate(doc):
                # 텍스트 추출
                page_text = page.get_text("text")
                
                if page_text.strip():
                    text_parts.append(page_text)
            
            # 텍스트가 충분히 추출되었으면 종료
            full_text = "\n\n".join(text_parts)
            
            if len(full_text.strip()) > 100:  # 의미있는 텍스트가 있으면
                doc.close()
                
                # pdfplumber로 테이블 감지
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            if page.extract_tables():
                                has_tables = True
                                break
                except:
                    pass
                
                # 텍스트 정제
                full_text = clean_pdf_text(full_text)
                
                # 메타데이터 추출
                metadata = self._extract_pdf_metadata(file_path)
                
                result = {
                    "text": full_text,
                    "page_count": page_count,
                    "has_tables": has_tables,
                    "confidence": 1.0,
                    "metadata": metadata
                }
                
                logger.info(f"PDF parsing completed: {page_count} pages, {len(full_text)} chars")
                return result
            
            # ---------------- 2. 이미지 기반 OCR (벡터 텍스트가 없는 경우) ----------------
            logger.info("Vector text not found, trying OCR...")
            reader = get_ocr_reader()
            ocr_text_parts = []
            
            for page_num in range(page_count):
                page = doc[page_num]
                # 고해상도로 렌더링 (3배 확대)
                pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
                
                # Pixmap을 NumPy 배열로 변환
                img_np = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
                    pix.height, pix.width, pix.n
                )
                
                # EasyOCR 실행
                results = reader.readtext(img_np, detail=0, paragraph=False)
                page_text = "\n".join(results)
                
                if page_text.strip():
                    ocr_text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            doc.close()
            
            if ocr_text_parts:
                full_text = "\n\n".join(ocr_text_parts)
                full_text = clean_pdf_text(full_text)
                
                result = {
                    "text": full_text,
                    "page_count": page_count,
                    "has_tables": False,
                    "confidence": 0.85,  # OCR 신뢰도
                    "metadata": self._extract_pdf_metadata(file_path)
                }
                
                logger.info(f"PDF OCR completed: {page_count} pages, {len(full_text)} chars")
                return result
            else:
                return {
                    "text": "[PDF OCR 결과 없음]",
                    "page_count": page_count,
                    "has_tables": False,
                    "confidence": 0.0,
                    "metadata": {}
                }
                
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            raise Exception(f"PDF 파싱 실패: {str(e)}")
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """PDF 메타데이터 추출"""
        try:
            reader = PdfReader(file_path)
            info = reader.metadata
            
            return {
                "title": info.get('/Title', ''),
                "author": info.get('/Author', ''),
                "subject": info.get('/Subject', ''),
                "creator": info.get('/Creator', ''),
                "producer": info.get('/Producer', ''),
                "creation_date": info.get('/CreationDate', ''),
            }
        except:
            return {}


class DOCXParser(DocumentParser):
    """DOCX 파일 파서 - 헤더/푸터 포함"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """DOCX 파싱 - 단락 + 표 + 헤더/푸터"""
        try:
            logger.info(f"Parsing DOCX file: {file_path}")
            
            doc = DocxDocument(file_path)
            text_parts = []
            has_tables = False
            
            # 본문 단락 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # 테이블 추출
            if doc.tables:
                has_tables = True
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text.strip())
            
            # 섹션별 헤더/푸터 추출
            for section in doc.sections:
                # 헤더
                header = section.header
                for p in header.paragraphs:
                    if p.text.strip():
                        text_parts.append(p.text.strip())
                
                # 푸터
                footer = section.footer
                for p in footer.paragraphs:
                    if p.text.strip():
                        text_parts.append(p.text.strip())
            
            # 전체 텍스트
            full_text = "\n\n".join(text_parts)
            full_text = self.clean_text(full_text)
            
            # 메타데이터
            metadata = {
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            }
            
            result = {
                "text": full_text if text_parts else "[DOCX 추출 결과 없음]",
                "page_count": len(doc.paragraphs),
                "has_tables": has_tables,
                "confidence": 1.0,
                "metadata": metadata
            }
            
            logger.info(f"DOCX parsing completed: {len(full_text)} chars")
            return result
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise Exception(f"DOCX 파싱 실패: {str(e)}")


class HWPParser(DocumentParser):
    """HWP 파일 파서 - HWPX 신버전 + OLE 구버전 지원"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """HWP 파싱 - 신버전(HWPX) 우선 → 구버전(OLE) 폴백"""
        try:
            logger.info(f"Parsing HWP file: {file_path}")
            
            file_extension = os.path.splitext(file_path)[1].lower()
            
            # ---------------- 1. 신버전 HWPX (ZIP/XML 기반) ----------------
            if file_extension == '.hwpx' or zipfile.is_zipfile(file_path):
                try:
                    text_content = []
                    
                    with zipfile.ZipFile(file_path, 'r') as z:
                        # HWPX 구조: Contents/*.xml에 텍스트 존재
                        for name in z.namelist():
                            if name.startswith('Contents/') and name.endswith('.xml'):
                                xml_data = z.read(name)
                                root = ET.fromstring(xml_data)
                                
                                # 네임스페이스 제거 (검색 단순화)
                                for elem in root.iter():
                                    if '}' in elem.tag:
                                        elem.tag = elem.tag.split('}', 1)[1]
                                
                                # 't' 태그에서 실제 텍스트 추출
                                for elem in root.iter('t'):
                                    if elem.text:
                                        text_content.append(elem.text)
                    
                    if text_content:
                        full_text = " ".join(text_content).strip()
                        full_text = self.clean_text(full_text)
                        
                        result = {
                            "text": full_text,
                            "page_count": 1,
                            "has_tables": False,
                            "confidence": 0.95,
                            "metadata": {}
                        }
                        
                        logger.info(f"HWPX parsing completed: {len(full_text)} chars")
                        return result
                        
                except Exception as e:
                    logger.warning(f"HWPX parsing failed, trying OLE format: {e}")
            
            # ---------------- 2. 구버전 OLE HWP ----------------
            if olefile.isOleFile(file_path):
                text_content = []
                
                with olefile.OleFileIO(file_path) as f:
                    # BodyText/SectionN 스트림 탐색
                    for i in range(256):
                        stream_name = f"BodyText/Section{i}"
                        if f.exists(stream_name):
                            try:
                                stream_data = f.openstream(stream_name).read()
                                
                                # UTF-16 디코딩
                                text = stream_data.decode('UTF-16', errors='ignore')
                                
                                # 제어 문자 제거
                                text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
                                
                                # Base64 같은 긴 문자열 제거
                                text = re.sub(r'[A-Za-z0-9+/=]{30,}', '', text)
                                
                                if text.strip():
                                    text_content.append(text.strip())
                                    
                            except Exception as e:
                                logger.warning(f"Failed to parse section {i}: {e}")
                                continue
                
                if text_content:
                    full_text = "\n\n".join(text_content)
                    full_text = self.clean_text(full_text)
                    
                    result = {
                        "text": full_text,
                        "page_count": len(text_content),
                        "has_tables": False,
                        "confidence": 0.9,
                        "metadata": {}
                    }
                    
                    logger.info(f"HWP OLE parsing completed: {len(full_text)} chars")
                    return result
                else:
                    return {
                        "text": "[HWP 추출 오류]: 본문 텍스트 스트림을 찾을 수 없습니다.",
                        "page_count": 0,
                        "has_tables": False,
                        "confidence": 0.0,
                        "metadata": {}
                    }
            
            # 지원하지 않는 형식
            return {
                "text": "[HWP 추출 오류]: 지원하지 않는 HWP 파일 형식입니다.",
                "page_count": 0,
                "has_tables": False,
                "confidence": 0.0,
                "metadata": {}
            }
            
        except Exception as e:
            logger.error(f"HWP parsing failed: {str(e)}")
            raise Exception(f"HWP 파싱 실패: {str(e)}")


class ImageParser(DocumentParser):
    """이미지 파일 파서 (EasyOCR) - Y좌표 정렬로 자연스러운 읽기 순서"""
    
    async def parse(self, file_path: str) -> Dict[str, Any]:
        """이미지 OCR - EasyOCR 사용"""
        try:
            logger.info(f"Parsing image file: {file_path}")
            
            # 이미지 파일 유효성 검사
            try:
                img = Image.open(file_path)
                img.verify()  # 이미지 손상 확인
                img = Image.open(file_path)  # verify() 후 다시 열기
            except Exception as img_error:
                logger.error(f"Image file read error: {img_error}")
                raise Exception(f"이미지 파일을 읽을 수 없습니다: {img_error}")
            
            # RGB로 변환
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            
            # NumPy 배열로 변환
            img_array = np.array(img)
            
            if img_array is None or img_array.size == 0:
                raise Exception("이미지 데이터가 비어 있습니다")
            
            logger.debug(f"Image loaded: {img_array.shape}")
            
            # EasyOCR 실행 (detail=1: bbox 포함)
            reader = get_ocr_reader()
            results = reader.readtext(img_array, detail=1, paragraph=False)
            
            if not results:
                return {
                    "text": "[이미지 OCR 결과 없음]",
                    "page_count": 1,
                    "has_tables": False,
                    "confidence": 0.0,
                    "metadata": {
                        "width": img.size[0],
                        "height": img.size[1],
                        "format": img.format,
                        "mode": img.mode,
                    }
                }
            
            # Y 좌표 기준으로 정렬 (위에서 아래로)
            sorted_results = sorted(results, key=lambda x: x[0][0][1])
            
            # 줄바꿈 감지 (Y 좌표 차이로 판단)
            lines = []
            current_line = []
            prev_y = None
            line_height_threshold = 30  # 픽셀 단위
            confidences = []
            
            for bbox, text, confidence in sorted_results:
                current_y = bbox[0][1]  # 좌상단 Y 좌표
                
                # 새로운 줄인지 판단
                if prev_y is None or abs(current_y - prev_y) > line_height_threshold:
                    if current_line:
                        lines.append(" ".join(current_line))
                    current_line = [text]
                else:
                    current_line.append(text)
                
                prev_y = current_y
                confidences.append(confidence)
            
            # 마지막 줄 추가
            if current_line:
                lines.append(" ".join(current_line))
            
            # 텍스트 정리
            full_text = "\n".join(lines)
            full_text = re.sub(r' +', ' ', full_text)  # 연속 공백 제거
            full_text = re.sub(r'\n\s+\n', '\n\n', full_text)  # 빈 줄 정리
            full_text = full_text.strip()
            
            # 평균 신뢰도 계산
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            # 이미지 메타데이터
            metadata = {
                "width": img.size[0],
                "height": img.size[1],
                "format": img.format,
                "mode": img.mode,
            }
            
            result = {
                "text": full_text,
                "page_count": 1,
                "has_tables": False,
                "confidence": avg_confidence,
                "metadata": metadata
            }
            
            logger.info(f"Image OCR completed: {len(full_text)} chars, confidence: {avg_confidence:.2%}")
            return result
            
        except Exception as e:
            logger.error(f"Image parsing failed: {str(e)}")
            raise Exception(f"이미지 파싱 실패: {str(e)}")


class DocumentParserFactory:
    """파서 팩토리 - 파일 타입에 따라 적절한 파서 선택"""
    
    # 파일 확장자별 파서 매핑
    PARSER_MAP = {
        '.pdf': PDFParser,
        '.docx': DOCXParser,
        '.doc': DOCXParser,
        '.hwp': HWPParser,
        '.jpg': ImageParser,
        '.jpeg': ImageParser,
        '.png': ImageParser,
        '.gif': ImageParser,
        '.bmp': ImageParser,
        '.tiff': ImageParser,
    }
    
    @classmethod
    def get_parser(cls, file_type: str) -> DocumentParser:
        """
        파일 타입에 맞는 파서 반환
        
        Args:
            file_type: 파일 확장자 (예: '.pdf', '.hwp')
            
        Returns:
            DocumentParser 인스턴스
            
        Raises:
            ValueError: 지원하지 않는 파일 타입
        """
        file_type = file_type.lower()
        
        parser_class = cls.PARSER_MAP.get(file_type)
        if not parser_class:
            supported = ', '.join(cls.PARSER_MAP.keys())
            raise ValueError(
                f"지원하지 않는 파일 형식입니다: {file_type}\n"
                f"지원 형식: {supported}"
            )
        
        return parser_class()
    
    @classmethod
    def detect_file_type(cls, file_path: str) -> str:
        """
        파일 경로에서 파일 타입 감지
        
        Args:
            file_path: 파일 경로
            
        Returns:
            파일 확장자 (예: '.pdf')
        """
        # 확장자 추출
        ext = Path(file_path).suffix.lower()
        
        # MIME 타입으로 검증
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # MIME 타입과 확장자가 일치하는지 확인
        if mime_type:
            logger.debug(f"Detected MIME type: {mime_type}")
        
        return ext
    
    @classmethod
    async def parse_file(cls, file_path: str) -> Dict[str, Any]:
        """
        파일을 자동으로 감지하여 파싱
        
        Args:
            file_path: 파싱할 파일 경로
            
        Returns:
            파싱 결과 딕셔너리
        """
        # 파일 존재 확인
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        # 파일 타입 감지
        file_type = cls.detect_file_type(file_path)
        
        # 파서 선택 및 파싱
        parser = cls.get_parser(file_type)
        result = await parser.parse(file_path)
        
        # 파일 타입 정보 추가
        result['file_type'] = file_type
        result['file_name'] = Path(file_path).name
        result['file_size'] = os.path.getsize(file_path)
        
        return result


# 편의 함수
async def parse_document(file_path: str) -> Dict[str, Any]:
    """
    문서 파싱 편의 함수
    
    Usage:
        result = await parse_document("document.pdf")
        print(result['text'])
    """
    return await DocumentParserFactory.parse_file(file_path)